"""Module for converting DOCX files to Markdown format."""

import logging
from pathlib import Path
from typing import List
try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

from .base import BaseConverter

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):
    """Converter for DOCX files to Markdown format."""
    
    def convert(self, file_path: Path, output_dir: Path) -> List[Path]:
        """
        Convert a DOCX file to Markdown format.
        
        Args:
            file_path (Path): Path to the DOCX file
            output_dir (Path): Directory to save the output files
            
        Returns:
            List[Path]: List of paths to the converted Markdown files
        """
        if not HAS_PYTHON_DOCX:
            raise ImportError("python-docx is required for DOCX conversion. Please install it with 'pip install python-docx'")
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        try:
            # Create output directory if it doesn't exist
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Load the DOCX document
            doc = Document(file_path)
            
            # Convert to Markdown
            md_content = self._docx_to_markdown(doc)
            
            # Save to file
            output_file = output_dir / f"{file_path.stem}.md"
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(md_content)
                
            logger.info(f"Converted {file_path.name} to {output_file}")
            return [output_file]
        except Exception as e:
            logger.error(f"Error converting DOCX file {file_path.name}: {e}")
            raise
    
    def _docx_to_markdown(self, doc) -> str:
        """Convert DOCX document to Markdown format."""
        md_lines = []
        
        for paragraph in doc.paragraphs:
            # Handle different heading styles
            if paragraph.style.name.startswith('Heading'):
                level = self._get_heading_level(paragraph.style.name)
                md_lines.append(f"{'#' * level} {paragraph.text}")
            else:
                # Handle text formatting
                text = self._format_text(paragraph)
                if text.strip():  # Only add non-empty lines
                    md_lines.append(text)
            
            # Add empty line after each paragraph
            md_lines.append("")
        
        # Handle tables
        for table in doc.tables:
            md_lines.extend(self._table_to_markdown(table))
            md_lines.append("")
            
        return "\n".join(md_lines)
    
    def _get_heading_level(self, style_name: str) -> int:
        """Get heading level from style name."""
        if 'Heading' in style_name:
            try:
                return int(style_name.replace('Heading ', ''))
            except ValueError:
                return 1
        return 1
    
    def _format_text(self, paragraph) -> str:
        """Format text with bold, italic, etc."""
        text_parts = []
        for run in paragraph.runs:
            text = run.text
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            text_parts.append(text)
        return "".join(text_parts)
    
    def _table_to_markdown(self, table) -> List[str]:
        """Convert table to Markdown format."""
        md_lines = []
        
        # Process header row
        if table.rows:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            md_lines.append("| " + " | ".join(header_cells) + " |")
            md_lines.append("|" + "|".join(["---"] * len(header_cells)) + "|")
            
            # Process data rows
            for row in table.rows[1:]:
                data_cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(data_cells) + " |")
                
        return md_lines
    
    def supported_extensions(self) -> List[str]:
        """
        Get list of supported file extensions.
        
        Returns:
            List[str]: List of supported extensions
        """
        return ['.docx']